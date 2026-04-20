import os
import tempfile
from typing import Optional
from sqlalchemy.orm import Session
from app.core.models import Document, FileType, Category
from app.services.ocr_service import extract_text_from_image
from app.services.text_processor import process_text, generate_summary, extract_keywords, categorize_document


async def process_file(file: any, db: Session) -> Document:
    """处理上传的文件"""
    # 获取文件扩展名
    filename = file.filename
    extension = os.path.splitext(filename)[1].lower()
    
    # 确定文件类型
    file_type = determine_file_type(extension)
    
    # 创建临时文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as temp_file:
        content = await file.read()
        temp_file.write(content)
        temp_file_path = temp_file.name
    
    try:
        # 提取文本
        text = extract_text(file_type, temp_file_path)
        
        # 处理文本
        processed_text = process_text(text)
        
        # 生成摘要
        summary = generate_summary(processed_text)
        
        # 提取关键词
        keywords = extract_keywords(processed_text)
        
        # 分类文档
        category = categorize_document(processed_text)
        
        # 创建文档记录
        document = Document(
            filename=filename,
            file_type=file_type,
            category=category,
            content=processed_text,
            summary=summary,
            keywords=keywords
        )
        
        # 保存到数据库
        db.add(document)
        db.commit()
        db.refresh(document)
        
        # 生成标签
        await generate_tags(document, processed_text, db)
        
        return document
    finally:
        # 清理临时文件
        if os.path.exists(temp_file_path):
            os.unlink(temp_file_path)


def determine_file_type(extension: str) -> FileType:
    """确定文件类型"""
    if extension in ['.doc', '.docx']:
        return FileType.WORD
    elif extension in ['.pdf']:
        return FileType.PDF
    elif extension in ['.jpg', '.jpeg', '.png', '.gif']:
        return FileType.IMAGE
    else:
        return FileType.OTHER


def extract_text(file_type: FileType, file_path: str) -> str:
    """提取文件中的文本"""
    if file_type == FileType.WORD:
        return extract_text_from_word(file_path)
    elif file_type == FileType.PDF:
        return extract_text_from_pdf(file_path)
    elif file_type == FileType.IMAGE:
        return extract_text_from_image(file_path)
    else:
        return ""


def extract_text_from_word(file_path: str) -> str:
    """从Word文件中提取文本"""
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return '\n'.join(text)


def extract_text_from_pdf(file_path: str) -> str:
    """从PDF文件中提取文本"""
    import PyPDF2
    text = []
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text.append(page.extract_text())
    return '\n'.join(text)


async def generate_tags(document: Document, text: str, db: Session):
    """为文档生成标签"""
    # 简单的标签生成逻辑
    # 实际应用中可以使用更复杂的NLP模型
    tag_names = set()
    
    # 从关键词中提取标签
    if document.keywords:
        keywords = document.keywords.split(',')
        for keyword in keywords:
            keyword = keyword.strip()
            if keyword and len(keyword) > 2:
                tag_names.add(keyword)
    
    # 保存标签
    for tag_name in tag_names:
        # 检查标签是否已存在
        tag = db.query(Tag).filter(Tag.name == tag_name).first()
        if not tag:
            tag = Tag(name=tag_name)
            db.add(tag)
            db.commit()
            db.refresh(tag)
        
        # 关联标签到文档
        if tag not in document.tags:
            document.tags.append(tag)
    
    db.commit()
    db.refresh(document)
